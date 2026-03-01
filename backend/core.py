import io
import json
import requests
import ast
import os
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

PROFILE_FIELD_ALIASES = {
    "å§“å": ["åå­—", "å§“åï¼ˆä¸­æ–‡ï¼‰", "å§“å(ä¸­æ–‡)", "name"],
    "æ€§åˆ«": ["gender"],
    "å‡ºç”Ÿæ—¥æœŸ": ["ç”Ÿæ—¥", "å‡ºç”Ÿå¹´æœˆ", "å‡ºç”Ÿæ—¶é—´", "birth", "dateofbirth"],
    "èº«ä»½è¯å·": ["èº«ä»½è¯", "èº«ä»½è¯å·ç ", "è¯ä»¶å·ç ", "id", "idcard"],
    "æ‰‹æœºå·ç ": ["æ‰‹æœºå·", "æ‰‹æœº", "ç”µè¯", "è”ç³»ç”µè¯", "è”ç³»æ–¹å¼", "mobile", "phone"],
    "ç”µå­é‚®ç®±": ["é‚®ç®±", "é‚®ä»¶", "email", "e-mail"],
    "æ¯•ä¸šé™¢æ ¡": ["æ¯•ä¸šå­¦æ ¡", "å­¦æ ¡", "é™¢æ ¡", "é«˜æ ¡", "university", "college"],
    "å­¦å†": ["æ•™è‚²ç¨‹åº¦", "education", "degree"],
    "ä¸“ä¸š": ["æ‰€å­¦ä¸“ä¸š", "major"],
    "æ¯•ä¸šæ—¶é—´": ["æ¯•ä¸šæ—¥æœŸ", "graduation", "graduationdate"],
    "åº”è˜å²—ä½": ["åº”è˜èŒä½", "æ±‚èŒå²—ä½", "èŒä½", "å²—ä½", "position", "jobtitle"],
    "æœŸæœ›åŸå¸‚": ["æ„å‘åŸå¸‚", "ç›®æ ‡åŸå¸‚", "æ±‚èŒåŸå¸‚", "expectedcity"],
    "ç°å±…ä½åœ°": ["ç°å±…", "å±…ä½åœ°", "å±…ä½åœ°å€", "åœ°å€", "address", "location"],
    "æ”¿æ²»é¢è²Œ": ["æ”¿æ²»èº«ä»½"],
    "ç´§æ€¥è”ç³»äºº": ["è”ç³»äºº", "ç´§æ€¥è”ç»œäºº", "emergencycontact"],
    "ç´§æ€¥è”ç³»äººç”µè¯": ["ç´§æ€¥è”ç³»äººæ‰‹æœºå·", "ç´§æ€¥è”ç³»ç”µè¯", "emergencyphone"],
}

PROFILE_KEY_NOISE_TOKENS = [
    "å¿…å¡«",
    "é€‰å¡«",
    "å¿…é€‰",
    "å¯é€‰",
    "required",
    "optional",
    "è¯·å¡«å†™",
    "è¯·è¾“å…¥",
]


def _normalize_profile_key(key):
    cleaned = re.sub(r"[\s_\-ï¼ˆï¼‰()ã€ã€‘\[\]Â·.]+", "", key or "")
    cleaned = cleaned.lower()
    for token in PROFILE_KEY_NOISE_TOKENS:
        cleaned = cleaned.replace(token, "")
    return cleaned.strip()


def _extract_profile_pairs(user_info_text):
    parsed_fields = {}
    for raw_line in (user_info_text or "").splitlines():
        line = raw_line.strip()
        if not line:
            continue

        line = re.sub(r"^[-*â€¢Â·]+\s*", "", line)
        line = re.sub(r"^\d+[.)ã€]\s*", "", line)

        if not re.search(r"[ï¼š:=]", line):
            continue

        parts = re.split(r"[ï¼š:=]", line, maxsplit=1)
        if len(parts) != 2:
            continue

        key = parts[0].strip()
        value = parts[1].strip().strip('"\'')
        if not key or not value:
            continue

        normalized_key = _normalize_profile_key(key)
        if normalized_key and normalized_key not in parsed_fields:
            parsed_fields[normalized_key] = value

    return parsed_fields


def _collect_explicit_profile_values(user_info_text):
    explicit_values = set()
    for value in _extract_profile_pairs(user_info_text).values():
        normalized_value = re.sub(r"\s+", "", str(value or ""))
        if normalized_value:
            explicit_values.add(normalized_value)
    return explicit_values


def _is_explicit_value(value, explicit_values, raw_text=""):
    normalized_value = re.sub(r"\s+", "", str(value or "")).strip()
    if not normalized_value:
        return False

    if normalized_value in explicit_values:
        return True

    # å…è®¸çŸ­è¯­åŒ…å«å…³ç³»ï¼ˆä¾‹å¦‚ "ä¸Šæµ·" vs "ä¸Šæµ·å¸‚æµ¦ä¸œæ–°åŒº"ï¼‰
    if len(normalized_value) <= 1:
        return False

    for explicit in explicit_values:
        if len(explicit) <= 1:
            continue
        if normalized_value in explicit or explicit in normalized_value:
            return True

    # å›é€€æ£€æŸ¥ï¼šåœ¨åŸå§‹ç”¨æˆ·æ–‡æœ¬ä¸­ç›´æ¥æœç´¢ï¼ˆå»é™¤ç©ºæ ¼ååŒ¹é…ï¼‰
    # è§£å†³ PDF æå–çš„æµå¼æ–‡æœ¬æ— æ³•è¢«è§£æä¸º key:value å¯¹çš„é—®é¢˜
    if raw_text:
        normalized_raw = re.sub(r"\s+", "", raw_text)
        if len(normalized_value) >= 2 and normalized_value in normalized_raw:
            return True

    return False


def build_profile_reuse_context(user_info_text):
    if not user_info_text:
        return user_info_text

    if "## æ ‡å‡†åŒ–èµ„æ–™æ˜ å°„ï¼ˆç³»ç»Ÿè‡ªåŠ¨ç”Ÿæˆï¼Œç”¨äºè·¨æ¨¡æ¿å¤ç”¨ï¼‰" in user_info_text:
        return user_info_text

    parsed_fields = _extract_profile_pairs(user_info_text)
    if not parsed_fields:
        return user_info_text

    canonical_pairs = []
    for canonical, aliases in PROFILE_FIELD_ALIASES.items():
        candidates = [canonical] + aliases
        matched_value = None
        for candidate in candidates:
            candidate_key = _normalize_profile_key(candidate)
            if candidate_key in parsed_fields:
                matched_value = parsed_fields[candidate_key]
                break

        if matched_value:
            canonical_pairs.append((canonical, matched_value))

    if not canonical_pairs:
        return user_info_text

    augmented_lines = ["", "## æ ‡å‡†åŒ–èµ„æ–™æ˜ å°„ï¼ˆç³»ç»Ÿè‡ªåŠ¨ç”Ÿæˆï¼Œç”¨äºè·¨æ¨¡æ¿å¤ç”¨ï¼‰"]
    augmented_lines.extend([f"{k}ï¼š{v}" for k, v in canonical_pairs])
    return user_info_text.rstrip() + "\n" + "\n".join(augmented_lines)

def analyze_missing_fields(docx_bytes, user_info_text):
    """
    åˆ†ææ¨¡æ¿å’Œä¸ªäººä¿¡æ¯ï¼Œè¿”å›å¯èƒ½ç¼ºå¤±çš„å­—æ®µåˆ—è¡¨

    Args:
        docx_bytes: Wordæ–‡æ¡£å­—èŠ‚æ•°æ®
        user_info_text: ç”¨æˆ·ä¿¡æ¯æ–‡æœ¬

    Returns:
        list: ç¼ºå¤±çš„å­—æ®µåç§°åˆ—è¡¨
    """
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))
    normalized_user_info_text = build_profile_reuse_context(user_info_text)

    # 1. æ”¶é›†è¡¨æ ¼çš„è¡¨å¤´å’Œå ä½ç¬¦ä¿¡æ¯
    placeholder_info = {}  # {å ä½ç¬¦: è¡¨æ ¼ä½ç½®æè¿°}
    all_headers = []  # æ‰€æœ‰è¡¨å¤´æ–‡æœ¬

    for t_idx, table in enumerate(doc.tables):
        if not table.rows:
            continue

        # è·å–è¡¨å¤´è¡Œ
        header_row = table.rows[0]
        headers = []
        for c_idx, cell in enumerate(header_row.cells):
            headers.append(cell.text.strip())

        # æ”¶é›†è¡¨å¤´ä¿¡æ¯
        for h_idx, header in enumerate(headers):
            if header and header not in all_headers:
                all_headers.append(header)

        # æ ‡è®°ç©ºå•å…ƒæ ¼å¹¶è®°å½•ä½ç½®
        max_cols = max(len(row.cells) for row in table.rows)
        counter = 1

        for r_idx, row in enumerate(table.rows):
            for c_idx in range(max_cols):
                if c_idx >= len(row.cells):
                    continue

                cell = row.cells[c_idx]
                text = cell.text.strip()

                if not text:
                    # ä¸ºç©ºå•å…ƒæ ¼åˆ›å»ºå ä½ç¬¦
                    tag = f"{{{counter}}}"
                    # å°è¯•æ‰¾åˆ°è¿™ä¸ªå ä½ç¬¦å¯¹åº”çš„è¡¨å¤´
                    header = headers[c_idx] if c_idx < len(headers) else ""
                    placeholder_info[tag] = {
                        "header": header,
                        "table_index": t_idx + 1,
                        "row_index": r_idx + 1,
                        "col_index": c_idx + 1
                    }
                    counter += 1

    if not placeholder_info:
        return []

    # 2. è°ƒç”¨ AI åˆ†æç¼ºå¤±çš„å­—æ®µ
    # å°†è¡¨æ ¼ä¿¡æ¯å’Œç”¨æˆ·ä¿¡æ¯ä¸€èµ·å‘ç»™ AIï¼Œè®©å®ƒæ¨æ–­éœ€è¦å“ªäº›å­—æ®µ
    headers_text = "\n".join([f"- {h}" for h in all_headers if h])
    placeholders_text = "\n".join([f"- {k}: è¡¨å¤´={v['header'] if v['header'] else 'æ— '}" for k, v in placeholder_info.items()])

    url = "https://api-inference.modelscope.cn/v1/chat/completions"
    api_key = os.environ.get("MODELSCOPE_API_KEY", "")
    model_endpoint = os.environ.get("MODEL_ENDPOINT") or "deepseek-ai/DeepSeek-V3.2"

    prompt = f"""ä½ æ˜¯ä¸€ä¸ªè¡¨å•å­—æ®µåˆ†æåŠ©æ‰‹ã€‚è¯·ä»”ç»†åˆ†æè¡¨æ ¼ä¸­çš„ç©ºå•å…ƒæ ¼å’Œä¸ªäººä¿¡æ¯ï¼Œæ‰¾å‡ºå“ªäº›å­—æ®µåœ¨ä¸ªäººä¿¡æ¯ä¸­æ²¡æœ‰æ˜ç¡®æä¾›ã€‚

**ä»»åŠ¡ï¼š**
1. æ‰¾å‡ºè¡¨æ ¼ä¸­æ‰€æœ‰ç©ºå•å…ƒæ ¼å¯¹åº”çš„å­—æ®µåç§°ï¼ˆæ ¹æ®è¡¨å¤´æˆ–åˆ—ä½ç½®åˆ¤æ–­ï¼‰
2. é€ä¸ªæ£€æŸ¥è¿™äº›å­—æ®µæ˜¯å¦åœ¨ã€ç”¨æˆ·å·²å¡«å†™çš„ä¿¡æ¯ã€‘ä¸­æœ‰æ˜ç¡®çš„å€¼
3. å¦‚æœæŸä¸ªå­—æ®µåœ¨ä¸ªäººä¿¡æ¯ä¸­æ²¡æœ‰æ˜ç¡®æä¾›ï¼Œå°±åŠ å…¥ç¼ºå¤±åˆ—è¡¨

**é‡è¦è§„åˆ™ï¼š**
- ä¸èƒ½é€šè¿‡æ¨ç†æˆ–çŒœæµ‹æ¥å¡«å……çš„å­—æ®µï¼Œå¿…é¡»æŠ¥å‘Šä¸ºç¼ºå¤±
- å³ä½¿å¯ä»¥é€šè¿‡å¸¸ç†æ¨æ–­çš„å­—æ®µï¼ˆå¦‚å®¶åº­æˆå‘˜çš„æ”¿æ²»é¢è²Œï¼‰ï¼Œå¦‚æœä¸ªäººä¿¡æ¯ä¸­æ²¡æœ‰æ˜ç¡®å†™å‡ºï¼Œä¹Ÿå¿…é¡»æŠ¥å‘Šä¸ºç¼ºå¤±
- ç…§ç‰‡ã€ç­¾åç­‰æ— æ³•é€šè¿‡æ–‡å­—æä¾›çš„å­—æ®µï¼Œä¹Ÿéœ€è¦æŠ¥å‘Š

**æ¨¡æ¿è¡¨æ ¼çš„ç©ºå•å…ƒæ ¼å­—æ®µï¼š**
{placeholders_text}

**ç”¨æˆ·å·²å¡«å†™çš„ä¿¡æ¯ï¼š**
{normalized_user_info_text}

**è¿”å›æ ¼å¼ï¼š**
è¯·ä»¥çº¯ JSON æ•°ç»„æ ¼å¼è¿”å›ï¼Œåªåˆ—å‡ºç¼ºå¤±çš„å­—æ®µåç§°ï¼Œç¤ºä¾‹ï¼š
["å°ä¸€å¯¸å½©è‰²è¿‘ç…§", "å®¶åº­æˆå‘˜æ”¿æ²»é¢è²Œ", "å®¶åº­æˆå‘˜è”ç³»ç”µè¯"]

åªè¿”å›å­—æ®µåç§°æ•°ç»„ï¼Œä¸è¦å…¶ä»–è§£é‡Šã€‚"""

    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    data = {
        "model": model_endpoint,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.3,  # è¾ƒä½æ¸©åº¦ä½¿ç»“æœæ›´ç¨³å®š
        "top_p": 0.7,
        "max_tokens": 500
    }

    try:
        response = requests.post(url, headers=headers, json=data)
        if response.status_code != 200:
            print(f"âŒ AI API è¿”å›é”™è¯¯çŠ¶æ€ç : {response.status_code}")
            return []

        res_json = response.json()
        content = res_json['choices'][0]['message']['content']

        print(f"ğŸ“ AI è¿”å›å†…å®¹: {content[:500]}...")

        # æ¸…ç† Markdown ä»£ç å—æ ‡è®°
        content = content.replace("```json", "").replace("```", "").strip()

        # è§£æ JSON æ•°ç»„
        try:
            missing_fields = json.loads(content)
            if isinstance(missing_fields, list):
                print(f"âœ… è§£æåˆ°ç¼ºå¤±å­—æ®µ: {missing_fields}")
                return missing_fields
            print(f"âš ï¸ è§£æç»“æœä¸æ˜¯æ•°ç»„: {missing_fields}")
            return []
        except json.JSONDecodeError as e:
            print(f"âš ï¸ JSON è§£æå¤±è´¥: {e}ï¼Œå°è¯•æ­£åˆ™æå–")
            # å°è¯•æå–æ•°ç»„æ ¼å¼
            import re
            matches = re.findall(r'"([^"]+)"', content)
            if matches:
                print(f"âœ… æ­£åˆ™æå–åˆ°ç¼ºå¤±å­—æ®µ: {matches}")
                return matches
            print(f"âš ï¸ æ­£åˆ™æå–å¤±è´¥ï¼ŒåŸå§‹å†…å®¹: {content[:200]}")
            return []
    except Exception as e:
        print(f"âŒ åˆ†æç¼ºå¤±å­—æ®µå¤±è´¥: {e}")
        return []


def audit_template(docx_bytes, user_info_text):
    """
    å®¡æ ¸æ¨¡æ¿å˜é‡ä¸ä¸ªäººä¿¡æ¯çš„åŒ¹é…æƒ…å†µ

    Args:
        docx_bytes: Wordæ–‡æ¡£å­—èŠ‚æ•°æ®
        user_info_text: ç”¨æˆ·ä¿¡æ¯æ–‡æœ¬

    Returns:
        dict: {
            "success": bool,
            "items": [{"key": str, "label": str, "value": str, "isMatched": bool}],
            "matched_count": int,
            "missing_count": int
        }
    """
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))
    normalized_user_info_text = build_profile_reuse_context(user_info_text)

    # 1. æ”¶é›†å ä½ç¬¦å’Œè¡¨å¤´ä¿¡æ¯ï¼Œæ„å»º Markdown è¡¨æ ¼
    placeholder_info = {}  # {å ä½ç¬¦: {"header": str, "table_index": int, "row_index": int, "col_index": int}}
    markdown_lines = []

    for t_idx, table in enumerate(doc.tables):
        if not table.rows:
            continue

        markdown_lines.append(f"\n### è¡¨æ ¼ {t_idx + 1}\n")

        # è·å–è¡¨å¤´è¡Œ
        header_row = table.rows[0]
        headers = []
        for c_idx, cell in enumerate(header_row.cells):
            headers.append(cell.text.strip())

        # é¢„è®¡ç®—æœ€å¤§åˆ—æ•°
        max_cols = max(len(row.cells) for row in table.rows)

        for r_idx, row in enumerate(table.rows):
            row_cells_content = []
            for c_idx in range(max_cols):
                if c_idx >= len(row.cells):
                    row_cells_content.append("")
                    continue

                cell = row.cells[c_idx]
                text = cell.text.strip()

                if not text:
                    # ç©ºå•å…ƒæ ¼ä½œä¸ºå ä½ç¬¦
                    tag = f"{{{len(placeholder_info) + 1}}}"
                    # è·å–è¡¨å¤´
                    header = headers[c_idx] if c_idx < len(headers) else ""
                    placeholder_info[tag] = {
                        "header": header,
                        "table_index": t_idx + 1,
                        "row_index": r_idx + 1,
                        "col_index": c_idx + 1
                    }
                    row_cells_content.append(tag)
                else:
                    row_cells_content.append(text)

            # ç”Ÿæˆ Markdown è¡Œ
            markdown_lines.append("| " + " | ".join(row_cells_content) + " |")
            if r_idx == 0:
                markdown_lines.append("| " + " | ".join(["---"] * max_cols) + " |")

    if not placeholder_info:
        return {"success": True, "items": [], "matched_count": 0, "missing_count": 0}

    # 2. è°ƒç”¨ AI åˆ†æåŒ¹é…æƒ…å†µ
    url = "https://api-inference.modelscope.cn/v1/chat/completions"
    api_key = os.environ.get("MODELSCOPE_API_KEY", "")
    model_endpoint = os.environ.get("MODEL_ENDPOINT") or "deepseek-ai/DeepSeek-V3.2"

    # æ„å»ºå ä½ç¬¦ä¿¡æ¯æ–‡æœ¬
    placeholders_text = "\n".join([
        f"- {k}: è¡¨å¤´=\"{v['header'] if v['header'] else 'æ— '}\" (è¡¨æ ¼{v['table_index']}ç¬¬{v['row_index']}è¡Œç¬¬{v['col_index']}åˆ—)"
        for k, v in placeholder_info.items()
    ])

    prompt = f"""ä½ æ˜¯ä¸€ä¸ªè¡¨å•åŒ¹é…å®¡æ ¸åŠ©æ‰‹ã€‚è¯·åˆ†æä»¥ä¸‹æ¨¡æ¿è¡¨æ ¼å’Œä¸ªäººä¿¡æ¯ï¼Œæ£€æŸ¥æ¯ä¸ªå ä½ç¬¦æ˜¯å¦èƒ½åœ¨ä¸ªäººä¿¡æ¯ä¸­æ‰¾åˆ°å¯¹åº”å€¼ã€‚

**ä»»åŠ¡ï¼š**
1. ä»æ¨¡æ¿è¡¨æ ¼ä¸­æå–æ‰€æœ‰å ä½ç¬¦åŠå…¶å«ä¹‰ï¼ˆæ ¹æ®è¡¨å¤´åˆ¤æ–­ï¼‰
2. åœ¨ç”¨æˆ·ä¿¡æ¯ä¸­æœç´¢æ¯ä¸ªå ä½ç¬¦å¯¹åº”çš„å€¼
3. å¦‚æœæ‰¾åˆ°å¯¹åº”å€¼ï¼Œæ ‡è®°ä¸ºåŒ¹é…ï¼›å¦‚æœæ‰¾ä¸åˆ°ï¼Œæ ‡è®°ä¸ºç¼ºå¤±

**æ¨¡æ¿è¡¨æ ¼çš„å ä½ç¬¦ä¿¡æ¯ï¼š**
{placeholders_text}

**Markdownè¡¨æ ¼ä¸Šä¸‹æ–‡ï¼š**
{"".join(markdown_lines)}

**ç”¨æˆ·å·²å¡«å†™çš„ä¿¡æ¯ï¼š**
{normalized_user_info_text}

**è¿”å›æ ¼å¼ï¼š**
è¯·ä»¥çº¯ JSON æ ¼å¼è¿”å›ï¼Œç¤ºä¾‹ï¼š
{{
  "items": [
    {{"key": "{{1}}", "label": "å§“å", "value": "å¼ ä¸‰", "isMatched": true}},
    {{"key": "{{2}}", "label": "æœŸæœ›è–ªèµ„", "value": "", "isMatched": false}}
  ]
}}

åªè¿”å› JSONï¼Œä¸è¦å…¶ä»–è§£é‡Šã€‚"""

    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    data = {
        "model": model_endpoint,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.3,
        "top_p": 0.7,
        "max_tokens": 2000
    }

    try:
        response = requests.post(url, headers=headers, json=data)
        if response.status_code != 200:
            print(f"âŒ AI API è¿”å›é”™è¯¯: {response.status_code}")
            return {"success": False, "error": f"API error: {response.status_code}", "items": []}

        res_json = response.json()
        content = res_json['choices'][0]['message']['content']

        # æ¸…ç† Markdown ä»£ç å—
        content = content.replace("```json", "").replace("```", "").strip()

        # è§£æ JSON
        result = json.loads(content)
        items = result.get("items", [])

        # ç¡®ä¿è¿”å›æ‰€æœ‰å ä½ç¬¦ï¼ˆAI å¯èƒ½é—æ¼ï¼‰
        returned_keys = {item.get("key") for item in items}
        for tag, info in placeholder_info.items():
            if tag not in returned_keys:
                items.append({
                    "key": tag,
                    "label": info["header"] if info["header"] else tag,
                    "value": "",
                    "isMatched": False
                })

        matched_count = sum(1 for item in items if item.get("isMatched"))
        missing_count = len(items) - matched_count

        return {
            "success": True,
            "items": items,
            "matched_count": matched_count,
            "missing_count": missing_count
        }

    except json.JSONDecodeError as e:
        print(f"âŒ JSON è§£æå¤±è´¥: {e}, å†…å®¹: {content[:200]}")
        return {"success": False, "error": f"JSON parse error: {str(e)}", "items": []}
    except Exception as e:
        print(f"âŒ å®¡æ ¸æ¨¡æ¿å¤±è´¥: {e}")
        return {"success": False, "error": str(e), "items": []}


def get_modelscope_response(user_info, markdown_context):
    """
    å‚è€ƒ smart.py çš„æç¤ºè¯æ€è·¯ï¼Œä½¿ç”¨ Markdown è¡¨æ ¼ä½œä¸ºä¸Šä¸‹æ–‡
    """
    if isinstance(user_info, bytes):
        user_info = user_info.decode('utf-8')

    url = "https://api-inference.modelscope.cn/v1/chat/completions"
    api_key = os.environ.get("MODELSCOPE_API_KEY", "")
    model_endpoint = os.environ.get("MODEL_ENDPOINT") or "deepseek-ai/DeepSeek-V3.2"

    # å‚è€ƒ smart.py çš„æç¤ºè¯æ„å»ºæ–¹å¼
    prompt = f"""ä½ æ˜¯ä¸€ä¸ªä¸“ä¸šçš„å ä½ç¬¦æ›¿æ¢åŠ©æ‰‹ã€‚è¯·åˆ†æä»¥ä¸‹ Markdown è¡¨æ ¼å’Œä¸ªäººä¿¡æ¯ï¼Œè¾“å‡ºæ¯ä¸ªå ä½ç¬¦åº”å¡«å†…å®¹ã€‚

**ä»»åŠ¡è¦æ±‚ï¼š**
1. ä»…åŸºäºã€ä¸ªäººä¿¡æ¯ã€‘ä¸­æ˜ç¡®å‡ºç°çš„å†…å®¹è¿›è¡Œå¡«å†™ï¼Œä¸å¾—ç¼–é€ ã€ä¸å¾—è„‘è¡¥ã€‚
2. ä»”ç»†åˆ†æè¡¨æ ¼ä¸­çš„å ä½ç¬¦æ ¼å¼ï¼ˆå¦‚ {{1}}ã€{{2}} ç­‰ï¼‰ï¼Œä»¥åŠå…¶æ‰€åœ¨è¡Œåˆ—ä¸Šä¸‹æ–‡ã€‚
3. å¦‚æœæ— æ³•ç¡®å®šæŸä¸ªå ä½ç¬¦çš„å†…å®¹ï¼Œå¿…é¡»è¿”å›ç©ºå­—ç¬¦ä¸² ""ã€‚
4. è¿”å›æ ¼å¼å¿…é¡»æ˜¯çº¯ JSONï¼Œæ ¼å¼ä¸ºï¼š{{"{{1}}": "å†…å®¹", "{{2}}": "å†…å®¹"}}ã€‚
5. è¿”å›å€¼å¿…é¡»æ˜¯ç®€æ´å­—æ®µå€¼ï¼Œä¸è¦è¾“å‡ºè§£é‡Šå¥ã€åŸå› æˆ–å¤šä½™å‰åç¼€ã€‚

**ä¸ªäººä¿¡æ¯ï¼š**
{user_info}

**Markdownè¡¨æ ¼ä¸Šä¸‹æ–‡ï¼š**
{markdown_context}

**æ³¨æ„ï¼š**
- åªè¿”å›éœ€è¦æ›¿æ¢çš„å ä½ç¬¦æ˜ å°„ã€‚
- ç¡®ä¿ JSON æ ¼å¼æ­£ç¡®ï¼Œä¸è¦åŒ…å«é¢å¤–çš„è§£é‡Šæ€§æ–‡å­—ã€‚"""

    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    data = {
        "model": model_endpoint, 
        "messages": [{"role": "user", "content": prompt}], 
        "temperature": 0.2,
        "top_p": 0.3,
        "extra_body": {"enable_thinking": True}
    }

    try:
        response = requests.post(url, headers=headers, json=data)
        if response.status_code != 200:
            return {}

        res_json = response.json()
        content = res_json['choices'][0]['message']['content']

        # æ¸…ç† Markdown ä»£ç å—æ ‡è®° (å‚è€ƒ smart.py çš„è§£æé€»è¾‘)
        content = content.replace("```json", "").replace("```", "").strip()

        # æ›´é²æ£’çš„ JSON æå–é€»è¾‘
        try:
            fill_data = json.loads(content)
        except json.JSONDecodeError:
            try:
                # å°è¯•åŒ¹é…ç¬¬ä¸€ä¸ª { å’Œæœ€åä¸€ä¸ª }
                match = re.search(r'\{.*\}', content, re.DOTALL)
                if match:
                    extracted_json = match.group(0)
                    print(f"ğŸ“ æå–åˆ° JSON: {extracted_json[:100]}...")
                    fill_data = json.loads(extracted_json)
                else:
                    print("âš ï¸ æœªæ‰¾åˆ° JSON æ ¼å¼å†…å®¹")
                    fill_data = {}
            except (json.JSONDecodeError, AttributeError) as e:
                print(f"âŒ JSON è§£æå¤±è´¥: {e}")
                try:
                    fill_data = ast.literal_eval(content)
                except:
                    print("âŒ æ‰€æœ‰ JSON è§£ææ–¹æ³•éƒ½å¤±è´¥")
                    fill_data = {}

        # ç»“æœå½’ä¸€åŒ–ï¼šä»…ä¿ç•™å ä½ç¬¦é”®ï¼Œå¹¶æ¸…ç†ç–‘ä¼¼è§£é‡Šæ€§æ–‡æœ¬
        normalized_fill_data = {}
        for key, value in (fill_data or {}).items():
            if not isinstance(key, str):
                continue

            normalized_key = key if key.startswith("{") else f"{{{key}}}"
            if not re.match(r"^\{\d+\}$", normalized_key):
                continue

            normalized_value = "" if value is None else str(value).strip()
            if any(token in normalized_value for token in ["æ— æ³•ç¡®å®š", "æœªæä¾›", "æœªçŸ¥", "æ ¹æ®æä¾›ä¿¡æ¯", "æ¨æ–­"]):
                normalized_value = ""

            normalized_fill_data[normalized_key] = normalized_value

        # æ‰“å° fill_data ä¾› server_with_auth.py è®°å½•
        print(f"ğŸ“‹ AI ç”Ÿæˆçš„å¡«å……æ•°æ®: {normalized_fill_data}")
        return normalized_fill_data
    except Exception as e:
        print(f"âŒ Error during AI inference: {e}")
        return {}

def _get_table_default_font(table):
    """è·å–è¡¨æ ¼çš„é»˜è®¤å­—ä½“æ ¼å¼ï¼ˆä»ç¬¬ä¸€ä¸ªæœ‰æ ¼å¼çš„éç©ºå•å…ƒæ ¼æå–ï¼‰"""
    default_font_name = None
    default_font_size = None
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if paragraph.runs:
                    run = paragraph.runs[0]
                    if run.font.name:
                        default_font_name = run.font.name
                    if run.font.size:
                        default_font_size = run.font.size
                    if default_font_name or default_font_size:
                        return default_font_name, default_font_size
    return default_font_name, default_font_size


def _replace_cell_text_preserve_format(cell, new_text, default_font_name=None, default_font_size=None):
    """
    æ›¿æ¢å•å…ƒæ ¼æ–‡æœ¬ï¼Œä¿æŒåŸæœ‰æ ¼å¼ã€‚
    å€Ÿé‰´ fill_template.py çš„æ€è·¯ï¼šä¿å­˜ run æ ¼å¼ â†’ åˆ é™¤æ—§ run â†’ åˆ›å»ºæ–° run â†’ æ¢å¤æ ¼å¼ã€‚
    """
    for paragraph in cell.paragraphs:
        # 1. ä¿å­˜ç¬¬ä¸€ä¸ª run çš„æ ¼å¼
        first_run_format = None
        if paragraph.runs:
            first_run = paragraph.runs[0]
            first_run_format = {
                'name': first_run.font.name,
                'size': first_run.font.size,
                'bold': first_run.font.bold,
                'italic': first_run.font.italic,
                'underline': first_run.font.underline,
            }

        # 2. åˆ é™¤æ‰€æœ‰ runï¼ˆä¿ç•™æ®µè½æœ¬èº«çš„å±æ€§å¦‚å¯¹é½æ–¹å¼ï¼‰
        for run in list(paragraph.runs):
            r = run._element
            r.getparent().remove(r)

        # 3. æ·»åŠ æ–° run å¹¶æ¢å¤æ ¼å¼
        new_run = paragraph.add_run(new_text)

        if first_run_format:
            if first_run_format['name']:
                new_run.font.name = first_run_format['name']
            if first_run_format['size']:
                new_run.font.size = first_run_format['size']
            if first_run_format['bold'] is not None:
                new_run.font.bold = first_run_format['bold']
            if first_run_format['italic'] is not None:
                new_run.font.italic = first_run_format['italic']
            if first_run_format['underline'] is not None:
                new_run.font.underline = first_run_format['underline']
        elif default_font_name or default_font_size:
            # æ²¡æœ‰åŸæ ¼å¼æ—¶ä½¿ç”¨è¡¨æ ¼é»˜è®¤æ ¼å¼
            if default_font_name:
                new_run.font.name = default_font_name
            if default_font_size:
                new_run.font.size = default_font_size

        # åªå¤„ç†ç¬¬ä¸€ä¸ªæ®µè½
        break


def fill_form(docx_bytes, user_info_text, photo_bytes, return_fill_data=False, prefilled_data=None, return_metadata=False):
    """
    å¡«å……è¡¨å•

    Args:
        docx_bytes: Wordæ–‡æ¡£å­—èŠ‚æ•°æ®
        user_info_text: ç”¨æˆ·ä¿¡æ¯æ–‡æœ¬
        photo_bytes: ç…§ç‰‡å­—èŠ‚æ•°æ®
        return_fill_data: æ˜¯å¦è¿”å›å¡«å……æ•°æ®ï¼ˆç”¨äºå‡å°‘é‡å¤æ¨ç†ï¼‰
        prefilled_data: å¯é€‰ï¼Œç›´æ¥ä½¿ç”¨é¢„è§ˆé˜¶æ®µè¿”å›çš„å¡«å……æ•°æ®ï¼Œé¿å…é‡å¤ AI æ¨ç†

    Returns:
        å¦‚æœ return_fill_data=Trueï¼Œè¿”å› (output_bytes, fill_data, missing_fields)
        å…¶ä¸­ missing_fields æ˜¯ç¼ºå¤±å­—æ®µçš„è¡¨å¤´/ä½ç½®ä¿¡æ¯åˆ—è¡¨
        å¦åˆ™è¿”å› output_bytes
    """
    doc = Document(io.BytesIO(docx_bytes))
    normalized_user_info_text = build_profile_reuse_context(user_info_text)
    explicit_profile_values = _collect_explicit_profile_values(normalized_user_info_text)

    # 1. å¤„ç†ç…§ç‰‡å ä½ç¬¦
    photo_coords = []
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                text_lower = cell.text.lower()
                if any(k in text_lower for k in ["ç…§ç‰‡", "ç›¸ç‰‡", "è¯ä»¶ç…§"]):
                    photo_coords.append((t_idx, r_idx, c_idx))

    if photo_coords and photo_bytes:
        for (t_idx, r_idx, c_idx) in photo_coords:
            cell = doc.tables[t_idx].rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(io.BytesIO(photo_bytes), width=Cm(3.5))

    # 2. æ ‡è®°ç©ºå•å…ƒæ ¼å¹¶æ„å»º Markdown ä¸Šä¸‹æ–‡ (é›†æˆ smart.py æ ¸å¿ƒæ€è·¯)
    placeholder_map = {}
    placeholder_info = {}  # å­˜å‚¨å ä½ç¬¦å¯¹åº”çš„è¡¨å¤´ä¿¡æ¯
    table_default_fonts = {}  # å­˜å‚¨æ¯ä¸ªè¡¨æ ¼çš„é»˜è®¤å­—ä½“ {t_idx: (font_name, font_size)}
    counter = 1
    markdown_lines = []

    for t_idx, table in enumerate(doc.tables):
        markdown_lines.append(f"\n### è¡¨æ ¼ {t_idx + 1}\n")

        # æå–è¡¨æ ¼é»˜è®¤å­—ä½“ï¼ˆç”¨äºç©ºå•å…ƒæ ¼çš„æ ¼å¼å›é€€ï¼‰
        table_default_fonts[t_idx] = _get_table_default_font(table)

        # é¢„è®¡ç®—å½“å‰è¡¨æ ¼çš„æœ€å¤§åˆ—æ•°
        max_cols = max(len(row.cells) for row in table.rows) if table.rows else 0

        for r_idx, row in enumerate(table.rows):
            row_cells_content = []
            for c_idx in range(max_cols):
                # è¶Šç•Œå¤„ç†
                if c_idx >= len(row.cells):
                    row_cells_content.append("")
                    continue

                cell = row.cells[c_idx]

                # è·³è¿‡å·²æ ‡è®°ä¸ºç…§ç‰‡çš„å•å…ƒæ ¼
                if (t_idx, r_idx, c_idx) in photo_coords:
                    row_cells_content.append("[ç…§ç‰‡]")
                    continue

                text = cell.text.strip()
                if not text:
                    # ä¸ºç©ºå•å…ƒæ ¼åˆ›å»ºå ä½ç¬¦
                    tag = f"{{{counter}}}"
                    # ä½¿ç”¨æ ¼å¼ä¿æŒçš„æ–¹å¼å†™å…¥å ä½ç¬¦æ ‡è®°
                    _replace_cell_text_preserve_format(
                        cell, tag,
                        table_default_fonts[t_idx][0],
                        table_default_fonts[t_idx][1]
                    )
                    placeholder_map[tag] = cell
                    # è·å–è¡¨å¤´ä¿¡æ¯
                    header = ""
                    if r_idx > 0 and c_idx < len(table.rows[0].cells):
                        header_cell = table.rows[0].cells[c_idx]
                        header = header_cell.text.strip()
                    # ä¿å­˜å ä½ç¬¦ä¿¡æ¯ï¼ŒåŒ…æ‹¬è¡¨å¤´å’Œä½ç½®
                    placeholder_info[tag] = {
                        "header": header,
                        "table_index": t_idx + 1,
                        "row_index": r_idx + 1,
                        "col_index": c_idx + 1
                    }
                    row_cells_content.append(tag)
                    counter += 1
                else:
                    row_cells_content.append(text)
            
            # ç”Ÿæˆ Markdown è¡Œ
            markdown_lines.append("| " + " | ".join(row_cells_content) + " |")
            if r_idx == 0: # æ·»åŠ åˆ†å‰²çº¿
                markdown_lines.append("| " + " | ".join(["---"] * max_cols) + " |")

    if not placeholder_map:
        out = io.BytesIO()
        doc.save(out)
        output_bytes = out.getvalue()
        if return_fill_data:
            if return_metadata:
                return output_bytes, {}, [], {"low_confidence_fields": []}
            return output_bytes, {}, []
        return output_bytes

    # 3. è·å–å¡«å……æ•°æ®ï¼ˆä¼˜å…ˆä½¿ç”¨é¢„è§ˆé˜¶æ®µä¼ å›çš„æ•°æ®ï¼Œé¿å…é‡å¤ AI æ¨ç†ï¼‰
    if prefilled_data is not None:
        fill_data = prefilled_data
    else:
        fill_data = get_modelscope_response(normalized_user_info_text, "\n".join(markdown_lines))

    if not isinstance(fill_data, dict):
        fill_data = {}

    # 4. æ”¶é›†æœªå¡«å……çš„å­—æ®µä¿¡æ¯
    missing_fields = []
    missing_fields_seen = set()
    placeholder_needs_ai_inference = {}
    resolved_placeholders = set()
    low_confidence_keys = set()

    def get_display_field_name(target_key, inferred_fields_map=None):
        header_info = placeholder_info.get(target_key, {})
        header = (header_info.get("header") or "").strip()
        if header:
            return header
        if inferred_fields_map and target_key in inferred_fields_map:
            candidate = str(inferred_fields_map[target_key]).strip()
            if candidate:
                return candidate
        return target_key

    def register_missing(target_key):
        header_info = placeholder_info.get(target_key, {})
        header = (header_info.get("header") or "").strip()
        if header:
            if header not in missing_fields_seen:
                missing_fields.append(header)
                missing_fields_seen.add(header)
        else:
            pos_info = placeholder_info.get(target_key)
            if pos_info and target_key not in placeholder_needs_ai_inference:
                placeholder_needs_ai_inference[target_key] = {
                    "table_index": pos_info.get("table_index", 0),
                    "row_index": pos_info.get("row_index", 0),
                    "col_index": pos_info.get("col_index", 0),
                }
        print(f"âš ï¸ è¯†åˆ«åˆ°ç¼ºå¤±å­—æ®µ: {header if header else target_key} (å ä½ç¬¦: {target_key})")

    # 4. å¡«å……æ•°æ®ï¼ˆä½¿ç”¨æ ¼å¼ä¿æŒçš„æ›¿æ¢æ–¹å¼ï¼‰
    for key, value in list(fill_data.items()):
        # å…¼å®¹ AI è¿”å› "1" è€Œä¸æ˜¯ "{1}" çš„æƒ…å†µ
        target_key = key if key.startswith("{") else f"{{{key}}}"
        if target_key in placeholder_map:
            resolved_placeholders.add(target_key)
            cell = placeholder_map[target_key]

            normalized_value = "" if value is None else str(value).strip()
            if normalized_value and not _is_explicit_value(normalized_value, explicit_profile_values, normalized_user_info_text):
                print(f"âš ï¸ ä½ç½®ä¿¡åº¦å€¼å·²æ¸…ç©º: {target_key} -> {normalized_value}")
                low_confidence_keys.add(target_key)
                normalized_value = ""

            fill_data[target_key] = normalized_value
            # æŸ¥æ‰¾è¯¥å•å…ƒæ ¼æ‰€å±è¡¨æ ¼çš„é»˜è®¤å­—ä½“
            cell_table_idx = placeholder_info.get(target_key, {}).get("table_index", 1) - 1
            def_font = table_default_fonts.get(cell_table_idx, (None, None))
            _replace_cell_text_preserve_format(cell, normalized_value, def_font[0], def_font[1])

            if not normalized_value:
                register_missing(target_key)

    # AI æœªè¿”å›æˆ–æœªå‘½ä¸­çš„å ä½ç¬¦ç»Ÿä¸€è§†ä¸ºç¼ºå¤±ï¼Œé¿å…ä¿ç•™ {n} æ ‡ç­¾
    for target_key, cell in placeholder_map.items():
        if target_key in resolved_placeholders:
            continue
        cell_table_idx = placeholder_info.get(target_key, {}).get("table_index", 1) - 1
        def_font = table_default_fonts.get(cell_table_idx, (None, None))
        _replace_cell_text_preserve_format(cell, "", def_font[0], def_font[1])
        fill_data[target_key] = ""
        register_missing(target_key)

    inferred_fields_map = {}

    # 5. å¦‚æœæœ‰æ— è¡¨å¤´çš„ç¼ºå¤±å­—æ®µï¼Œç”¨ AI æ¨æ–­å­—æ®µåç§°
    if placeholder_needs_ai_inference:
        placeholder_keys = list(placeholder_needs_ai_inference.keys())
        inferred_fields = infer_field_names_with_ai(
            placeholder_needs_ai_inference,
            "\n".join(markdown_lines),
            normalized_user_info_text
        )
        inferred_fields_map = {
            key: inferred_fields[idx]
            for idx, key in enumerate(placeholder_keys)
            if idx < len(inferred_fields)
        }
        for field in inferred_fields:
            candidate = str(field).strip()
            if candidate and candidate not in missing_fields_seen:
                missing_fields.append(candidate)
                missing_fields_seen.add(candidate)

    low_confidence_fields = []
    low_confidence_seen = set()
    for target_key in low_confidence_keys:
        display_name = get_display_field_name(target_key, inferred_fields_map)
        if display_name and display_name not in low_confidence_seen:
            low_confidence_fields.append(display_name)
            low_confidence_seen.add(display_name)

    print(f"ğŸ“‹ ç¼ºå¤±å­—æ®µåˆ—è¡¨: {missing_fields}")
    if low_confidence_fields:
        print(f"ğŸ“‰ ä½ç½®ä¿¡åº¦å­—æ®µåˆ—è¡¨: {low_confidence_fields}")

    out = io.BytesIO()
    doc.save(out)
    output_bytes = out.getvalue()

    if return_fill_data:
        if return_metadata:
            return output_bytes, fill_data, missing_fields, {
                "low_confidence_fields": low_confidence_fields
            }
        return output_bytes, fill_data, missing_fields
    return output_bytes


def infer_field_names_with_ai(placeholder_info_map, markdown_context, user_info_text):
    """
    ä½¿ç”¨ AI æ¨æ–­ç¼ºå¤±å­—æ®µçš„åç§°ï¼ˆå½“è¡¨å¤´ä¸ºç©ºæ—¶ï¼‰

    Args:
        placeholder_info_map: å ä½ç¬¦ä¿¡æ¯å­—å…¸ {å ä½ç¬¦: {table_index, row_index, col_index}}
        markdown_context: Markdown è¡¨æ ¼ä¸Šä¸‹æ–‡
        user_info_text: ç”¨æˆ·å·²å¡«å†™çš„ä¿¡æ¯

    Returns:
        list: æ¨æ–­å‡ºçš„å­—æ®µåç§°åˆ—è¡¨
    """
    if not placeholder_info_map:
        return []

    url = "https://api-inference.modelscope.cn/v1/chat/completions"
    api_key = os.environ.get("MODELSCOPE_API_KEY", "")
    model_endpoint = os.environ.get("MODEL_ENDPOINT") or "deepseek-ai/DeepSeek-V3.2"

    # æ„å»ºå ä½ç¬¦ä¿¡æ¯
    placeholders_text = "\n".join([
        f"- {k}: è¡¨æ ¼{v['table_index']}ç¬¬{v['row_index']}è¡Œç¬¬{v['col_index']}åˆ—"
        for k, v in placeholder_info_map.items()
    ])

    prompt = f"""ä½ æ˜¯ä¸€ä¸ªè¡¨å•å­—æ®µåˆ†æåŠ©æ‰‹ã€‚è¡¨æ ¼ä¸­æœ‰ä¸€äº›ç©ºå•å…ƒæ ¼æ²¡æœ‰è¡¨å¤´ï¼Œéœ€è¦ä½ æ ¹æ®ä¸Šä¸‹æ–‡æ¨æ–­è¿™äº›å•å…ƒæ ¼åº”è¯¥å¡«å†™ä»€ä¹ˆç±»å‹çš„å­—æ®µã€‚

**ä»»åŠ¡ï¼š**
åˆ†æè¡¨æ ¼ç»“æ„å’Œç”¨æˆ·ä¿¡æ¯ï¼Œæ¨æ–­æ¯ä¸ªç©ºå•å…ƒæ ¼åº”è¯¥å¡«å†™ä»€ä¹ˆç±»å‹çš„å­—æ®µåç§°ï¼ˆå¦‚"èº«é«˜"ã€"ä½“é‡"ã€"æ¯•ä¸šé™¢æ ¡"ç­‰ï¼‰ã€‚

**è¡¨æ ¼ä¸Šä¸‹æ–‡ï¼š**
{markdown_context}

**éœ€è¦æ¨æ–­çš„å ä½ç¬¦ä½ç½®ï¼š**
{placeholders_text}

**ç”¨æˆ·å·²å¡«å†™çš„ä¿¡æ¯ï¼š**
{user_info_text}

**è¿”å›æ ¼å¼ï¼š**
è¯·ä»¥çº¯ JSON æ•°ç»„æ ¼å¼è¿”å›å­—æ®µåç§°ï¼Œé¡ºåºä¸å ä½ç¬¦é¡ºåºä¸€è‡´ï¼Œç¤ºä¾‹ï¼š
["èº«é«˜(cm)", "ä½“é‡(kg)", "æ¯•ä¸šé™¢æ ¡"]

åªè¿”å›å­—æ®µåç§°ï¼Œä¸è¦å…¶ä»–è§£é‡Šã€‚å¦‚æœæ²¡æœ‰è¶³å¤Ÿä¿¡æ¯æ¨æ–­ï¼Œå¯ä»¥ä½¿ç”¨é€šç”¨æè¿°å¦‚"å­—æ®µ"ã€"ä¿¡æ¯"ç­‰ã€‚"""

    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    data = {
        "model": model_endpoint,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.3,
        "top_p": 0.7,
        "max_tokens": 500
    }

    try:
        response = requests.post(url, headers=headers, json=data)
        if response.status_code != 200:
            # å¦‚æœ AI è°ƒç”¨å¤±è´¥ï¼Œè¿”å›å ä½ç¬¦ä½œä¸ºé»˜è®¤
            return list(placeholder_info_map.keys())

        res_json = response.json()
        content = res_json['choices'][0]['message']['content']

        # æ¸…ç† Markdown ä»£ç å—æ ‡è®°
        content = content.replace("```json", "").replace("```", "").strip()

        # è§£æ JSON æ•°ç»„
        try:
            inferred_fields = json.loads(content)
            if isinstance(inferred_fields, list):
                # ç¡®ä¿è¿”å›æ•°é‡ä¸å ä½ç¬¦æ•°é‡ä¸€è‡´
                if len(inferred_fields) == len(placeholder_info_map):
                    return inferred_fields
                else:
                    # æ•°é‡ä¸åŒ¹é…æ—¶ï¼Œè¡¥å……æˆ–æˆªæ–­
                    placeholders = list(placeholder_info_map.keys())
                    if len(inferred_fields) < len(placeholders):
                        return inferred_fields + placeholders[len(inferred_fields):]
                    else:
                        return inferred_fields[:len(placeholders)]
            return list(placeholder_info_map.keys())
        except json.JSONDecodeError:
            # å°è¯•æå–æ•°ç»„æ ¼å¼
            matches = re.findall(r'"([^"]+)"', content)
            if matches:
                return matches
            return list(placeholder_info_map.keys())
    except Exception as e:
        print(f"âŒ AI æ¨æ–­å­—æ®µåç§°å¤±è´¥: {e}")
        return list(placeholder_info_map.keys())