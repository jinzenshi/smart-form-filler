#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
改进的报名表填写工具
基于 smart_replace_placeholders.py 的高效方法
同时保留照片处理功能
"""

import io
import json
import os
import re
import requests
from typing import Dict, Tuple, List
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm


class FormFiller:
    """智能表单填写器"""

    def __init__(self):
        self.api_key = os.environ.get("ARK_API_KEY") or "5410d463-1115-4320-9279-a5441ce30694"
        self.model_endpoint = os.environ.get("MODEL_ENDPOINT") or "doubao-seed-1-6-251015"
        self.url = "https://ark.cn-beijing.volces.com/api/v3/chat/completions"

    def convert_to_markdown(self, doc: Document) -> Tuple[str, List[Tuple]]:
        """
        将docx转换为markdown，返回markdown内容和照片坐标
        照片坐标格式: (table_idx, row_idx, cell_idx)
        """
        markdown_lines = []
        photo_coords = []

        # 首先识别照片位置
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    text_lower = cell.text.lower()
                    if ("照片" in text_lower) or ("相片" in text_lower) or ("证件照" in text_lower):
                        photo_coords.append((t_idx, r_idx, c_idx))

        # 转换为markdown
        for table_idx, table in enumerate(doc.tables):
            markdown_lines.append(f"\n## 表格 {table_idx + 1}\n")

            # 获取表格最大行列数
            max_rows = len(table.rows)
            max_cols = max(len(row.cells) for row in table.rows)

            # 生成markdown表格
            for row_idx in range(max_rows):
                row_cells = []
                for col_idx in range(max_cols):
                    if row_idx < len(table.rows) and col_idx < len(table.rows[row_idx].cells):
                        cell = table.rows[row_idx].cells[col_idx]

                        # 检查是否是照片位置
                        if (table_idx, row_idx, col_idx) in photo_coords:
                            cell_text = "[照片]"
                        else:
                            cell_text = cell.text.strip()
                            if not cell_text:
                                cell_text = "<empty>"
                    else:
                        cell_text = "<empty>"

                    row_cells.append(cell_text)

                markdown_lines.append("| " + " | ".join(row_cells) + " |")

                if row_idx == 0:
                    separator = "| " + " | ".join(["---"] * len(row_cells)) + " |"
                    markdown_lines.append(separator)

        return '\n'.join(markdown_lines), photo_coords

    def insert_photos(self, doc: Document, photo_coords: List[Tuple], photo_bytes: bytes):
        """在指定位置插入照片"""
        if not photo_coords or not photo_bytes:
            return

        for (t_idx, r_idx, c_idx) in photo_coords:
            cell = doc.tables[t_idx].rows[r_idx].cells[c_idx]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(io.BytesIO(photo_bytes), width=Cm(3.5))

    def call_ai_to_fill(self, user_info: str, markdown_content: str) -> Dict[str, str]:
        """调用AI分析占位符并返回填充数据"""
        print(f"\n🔑 API Key: {self.api_key[:10]}...")
        print(f"🤖 Model: {self.model_endpoint}")
        print(f"📝 User Info Length: {len(user_info)}")
        print(f"📋 Markdown Content Length: {len(markdown_content)}")

        prompt = f"""
        你是一个智能填表助手。

        【任务】
        表格中的空缺项已标记为 <empty> 或占位符如 {{1}}, {{2}}...
        请根据【个人资料】推断每个空项应该填写的内容。

        【个人资料】
        {user_info}

        【表格上下文】
        {markdown_content}

        【要求】
        1. 返回纯 JSON，格式 {{"{{1}}": "内容", "{{2}}": "内容"}}。
        2. 找不到信息填 "无"。
        3. 如果表格中没有占位符，返回空对象 {{}}。
        4. 允许合理推理和推断。
        """

        headers = {"Authorization": f"Bearer {self.api_key}", "Content-Type": "application/json"}
        data = {
            "model": self.model_endpoint,
            "messages": [{"role": "user", "content": prompt}],
            "thinking": {"type": "disabled"},
            "top_p": 0.7,
            "temperature": 1
        }

        try:
            response = requests.post(self.url, headers=headers, json=data)
            print(f"📡 Response Status: {response.status_code}")

            if response.status_code != 200:
                print(f"❌ Error Response: {response.text}")
                return {}

            res_json = response.json()
            print(f"🔍 Response Keys: {res_json.keys()}")

            if 'choices' not in res_json or not res_json['choices']:
                print("❌ No choices in response")
                return {}

            content = res_json['choices'][0]['message']['content']
            # 清理可能的代码块标记
            content = content.replace("```json", "").replace("```", "").strip()
            print(f"📄 Raw Content: {content[:200]}...")

            try:
                result = json.loads(content)
                print(f"✅ Parsed JSON: {result}")
                return result
            except json.JSONDecodeError as e:
                print(f"⚠️ JSON decode failed: {e}")
                try:
                    result = ast.literal_eval(content)
                    print(f"✅ Parsed with ast: {result}")
                    return result
                except Exception as e2:
                    print(f"❌ AST parse failed: {e2}")
                    return {}

        except Exception as e:
            print(f"❌ Exception: {e}")
            import traceback
            traceback.print_exc()
            return {}

    def replace_placeholders(self, doc: Document, fill_data: Dict[str, str]):
        """替换docx中的占位符"""
        total_replaced = 0

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    original_text = cell.text
                    new_text = original_text

                    # 替换各种格式的占位符
                    for key, value in fill_data.items():
                        # 处理 {1}, {2} 格式
                        placeholder_patterns = [
                            f"{{{key}}}",  # {1}
                            key,           # 1
                            f"<{key}>",    # <1>
                        ]

                        for pattern in placeholder_patterns:
                            if pattern in new_text:
                                new_text = new_text.replace(pattern, str(value))
                                total_replaced += 1

                    # 如果内容有变化，更新单元格
                    if new_text != original_text:
                        cell.text = new_text
                        # 设置左对齐
                        for p in cell.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        print(f"✅ 共替换了 {total_replaced} 个占位符")
        return total_replaced

    def save_replacement_map(self, fill_data: Dict[str, str], output_path: str):
        """保存填充映射到JSON文件"""
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(fill_data, f, ensure_ascii=False, indent=2)
        print(f"📄 填充映射已保存到: {output_path}")

    def fill_form(self, docx_bytes: bytes, user_info_text: str, photo_bytes: bytes = None,
                  output_json_path: str = None) -> bytes:
        """
        完整的表单填写流程

        Args:
            docx_bytes: docx文件的字节数据
            user_info_text: 个人信息的文本
            photo_bytes: 可选的照片字节数据
            output_json_path: 可选的JSON输出路径

        Returns:
            填写完成的docx字节数据
        """
        print("=" * 60)
        print("智能表单填写程序")
        print("=" * 60)

        # 1. 加载文档
        print("\n【步骤1】加载文档...")
        doc = Document(io.BytesIO(docx_bytes))
        print("✅ 文档加载完成")

        # 2. 转换为markdown
        print("\n【步骤2】转换为Markdown...")
        markdown_content, photo_coords = self.convert_to_markdown(doc)
        print("✅ 转换完成")

        # 3. 保存markdown（可选）
        markdown_path = "form_markdown.md"
        with open(markdown_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        print(f"📝 Markdown已保存到: {markdown_path}")

        # 4. 插入照片（如果提供）
        if photo_bytes:
            print("\n【步骤3】插入照片...")
            self.insert_photos(doc, photo_coords, photo_bytes)
            print("✅ 照片插入完成")

        # 5. 调用AI填充
        print("\n【步骤4】调用AI分析并填充...")
        fill_data = self.call_ai_to_fill(user_info_text, markdown_content)

        if not fill_data:
            print("⚠️ 没有获取到填充数据，返回原始文档")
            out = io.BytesIO()
            doc.save(out)
            return out.getvalue()

        # 6. 保存填充映射
        if output_json_path:
            print("\n【步骤5】保存填充映射...")
            self.save_replacement_map(fill_data, output_json_path)

        # 7. 替换占位符
        print("\n【步骤6】替换占位符...")
        self.replace_placeholders(doc, fill_data)

        # 8. 保存结果
        print("\n【步骤7】保存文档...")
        out = io.BytesIO()
        doc.save(out)
        result_bytes = out.getvalue()
        print("✅ 文档保存完成")

        print("\n" + "=" * 60)
        print("✅ 所有步骤完成！")
        print("=" * 60)

        return result_bytes


def fill_form(docx_bytes, user_info_text, photo_bytes=None):
    """
    便捷函数：填写表单
    保持与原core.py的兼容性

    Args:
        docx_bytes: docx文件的字节数据
        user_info_text: 个人信息的文本
        photo_bytes: 可选的照片字节数据

    Returns:
        填写完成的docx字节数据
    """
    filler = FormFiller()
    return filler.fill_form(docx_bytes, user_info_text, photo_bytes)


# 如果直接运行此脚本，提供示例用法
if __name__ == "__main__":
    # 示例用法
    import sys

    if len(sys.argv) < 3:
        print("用法: python core_improved.py <docx文件> <个人信息文件> [照片文件]")
        sys.exit(1)

    docx_file = sys.argv[1]
    info_file = sys.argv[2]
    photo_file = sys.argv[3] if len(sys.argv) > 3 else None

    # 读取文件
    with open(docx_file, 'rb') as f:
        docx_bytes = f.read()

    with open(info_file, 'r', encoding='utf-8') as f:
        user_info = f.read()

    photo_bytes = None
    if photo_file and os.path.exists(photo_file):
        with open(photo_file, 'rb') as f:
            photo_bytes = f.read()

    # 填写表单
    filler = FormFiller()
    result = filler.fill_form(docx_bytes, user_info, photo_bytes, "fill_result.json")

    # 保存结果
    output_file = docx_file.replace('.docx', '_filled.docx')
    with open(output_file, 'wb') as f:
        f.write(result)

    print(f"\n🎉 完成！结果已保存到: {output_file}")
