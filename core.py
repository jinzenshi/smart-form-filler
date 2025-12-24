import io
import json
import requests
import ast
import os
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

def get_doubao_response(user_info, markdown_context):
    """
    参考 smart.py 的提示词思路，使用 Markdown 表格作为上下文
    """
    if isinstance(user_info, bytes):
        user_info = user_info.decode('utf-8')

    url = "https://ark.cn-beijing.volces.com/api/v3/chat/completions"
    api_key = os.environ.get("ARK_API_KEY") or "5410d463-1115-4320-9279-a5441ce30694"
    model_endpoint = os.environ.get("MODEL_ENDPOINT") or "doubao-seed-1-6-251015"

    # 参考 smart.py 的提示词构建方式
    prompt = f"""你是一个专业的占位符替换助手。请分析以下 Markdown 表格和个人信息，推断每个占位符应该替换成什么内容。

**任务要求：**
1. 仔细分析表格中的占位符格式（如 {{1}}、{{2}} 等）。
2. 根据【个人信息】推理每个占位符应该填入的内容，允许合理推理和推断。
3. 如果无法确定某个占位符的内容，返回空字符串。
4. 返回格式必须是纯 JSON，格式为：{{"{{1}}": "内容", "{{2}}": "内容"}}。
5. 文字过长请注意换行来让排版更美观。

**个人信息：**
{user_info}

**Markdown表格上下文：**
{markdown_context}

**注意：**
- 只返回需要替换的占位符映射。
- 确保 JSON 格式正确，不要包含额外的解释性文字。"""

    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    data = {
        "model": model_endpoint, 
        "messages": [{"role": "user", "content": prompt}], 
        "temperature": 1, 
        "top_p": 0.7,
        "thinking": {"type": "disabled"}
    }

    try:
        response = requests.post(url, headers=headers, json=data)
        if response.status_code != 200:
            return {}

        res_json = response.json()
        content = res_json['choices'][0]['message']['content']

        # 清理 Markdown 代码块标记 (参考 smart.py 的解析逻辑)
        content = content.replace("```json", "").replace("```", "").strip()

        # 更鲁棒的 JSON 提取逻辑
        try:
            fill_data = json.loads(content)
        except json.JSONDecodeError:
            try:
                # 尝试匹配第一个 { 和最后一个 }
                match = re.search(r'\{.*\}', content, re.DOTALL)
                if match:
                    extracted_json = match.group(0)
                    print(f"📝 提取到 JSON: {extracted_json[:100]}...")
                    fill_data = json.loads(extracted_json)
                else:
                    print("⚠️ 未找到 JSON 格式内容")
                    fill_data = {}
            except (json.JSONDecodeError, AttributeError) as e:
                print(f"❌ JSON 解析失败: {e}")
                try:
                    fill_data = ast.literal_eval(content)
                except:
                    print("❌ 所有 JSON 解析方法都失败")
                    fill_data = {}

        # 打印 fill_data 供 server_with_auth.py 记录
        print(f"📋 AI 生成的填充数据: {fill_data}")
        return fill_data
    except Exception as e:
        print(f"❌ Error during AI inference: {e}")
        return {}

def fill_form(docx_bytes, user_info_text, photo_bytes, return_fill_data=False):
    """
    填充表单

    Args:
        docx_bytes: Word文档字节数据
        user_info_text: 用户信息文本
        photo_bytes: 照片字节数据
        return_fill_data: 是否返回填充数据（用于减少重复推理）

    Returns:
        如果 return_fill_data=True，返回 (output_bytes, fill_data)
        否则返回 output_bytes
    """
    doc = Document(io.BytesIO(docx_bytes))

    # 1. 处理照片占位符
    photo_coords = []
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                text_lower = cell.text.lower()
                if any(k in text_lower for k in ["照片", "相片", "证件照"]):
                    photo_coords.append((t_idx, r_idx, c_idx))

    if photo_coords and photo_bytes:
        for (t_idx, r_idx, c_idx) in photo_coords:
            cell = doc.tables[t_idx].rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(io.BytesIO(photo_bytes), width=Cm(3.5))

    # 2. 标记空单元格并构建 Markdown 上下文 (集成 smart.py 核心思路)
    placeholder_map = {}
    counter = 1
    markdown_lines = []

    for t_idx, table in enumerate(doc.tables):
        markdown_lines.append(f"\n### 表格 {t_idx + 1}\n")
        
        # 预计算当前表格的最大列数
        max_cols = max(len(row.cells) for row in table.rows) if table.rows else 0
        
        for r_idx, row in enumerate(table.rows):
            row_cells_content = []
            for c_idx in range(max_cols):
                # 越界处理
                if c_idx >= len(row.cells):
                    row_cells_content.append("")
                    continue
                
                cell = row.cells[c_idx]
                
                # 跳过已标记为照片的单元格
                if (t_idx, r_idx, c_idx) in photo_coords:
                    row_cells_content.append("[照片]")
                    continue
                
                text = cell.text.strip()
                if not text:
                    # 为空单元格创建占位符
                    tag = f"{{{counter}}}"
                    cell.text = tag
                    placeholder_map[tag] = cell
                    row_cells_content.append(tag)
                    counter += 1
                else:
                    row_cells_content.append(text)
            
            # 生成 Markdown 行
            markdown_lines.append("| " + " | ".join(row_cells_content) + " |")
            if r_idx == 0: # 添加分割线
                markdown_lines.append("| " + " | ".join(["---"] * max_cols) + " |")

    if not placeholder_map:
        out = io.BytesIO()
        doc.save(out)
        output_bytes = out.getvalue()
        if return_fill_data:
            return output_bytes, {}
        return output_bytes

    # 3. 调用 AI 进行推理
    fill_data = get_doubao_response(user_info_text, "\n".join(markdown_lines))

    # 4. 填充数据
    if fill_data:
        for key, value in fill_data.items():
            # 兼容 AI 返回 "1" 而不是 "{1}" 的情况
            target_key = key if key.startswith("{") else f"{{{key}}}"
            if target_key in placeholder_map:
                cell = placeholder_map[target_key]
                cell.text = str(value)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    out = io.BytesIO()
    doc.save(out)
    output_bytes = out.getvalue()

    if return_fill_data:
        return output_bytes, fill_data
    return output_bytes