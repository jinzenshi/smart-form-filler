#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
智能替换docx中的占位符
1. 将带占位符的docx转换为markdown
2. 将markdown和个人信息发送给AI
3. AI返回JSON格式的替换映射
4. 根据JSON替换docx中的占位符
5. 保存新的docx文件
"""

import os
import re
import json
import requests
from typing import Dict, Tuple
from docx import Document


class PlaceholderReplacer:
    """占位符智能替换器"""

    def __init__(self, docx_path: str, personal_info_path: str, api_key: str):
        self.docx_path = docx_path
        self.personal_info_path = personal_info_path
        self.api_key = api_key
        self.document = Document(docx_path)
        self.personal_info = self._read_personal_info()

    def _read_personal_info(self) -> str:
        """读取完整个人信息"""
        with open(self.personal_info_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return content

    def convert_to_markdown(self) -> str:
        """将docx转换为markdown"""
        markdown_lines = []

        for table_idx, table in enumerate(self.document.tables):
            markdown_lines.append(f"\n## 表格 {table_idx + 1}\n")

            # 获取表格最大行列数
            max_rows = len(table.rows)
            max_cols = max(len(row.cells) for row in table.rows)

            # 生成markdown
            for row_idx in range(max_rows):
                row_cells = []
                for col_idx in range(max_cols):
                    if row_idx < len(table.rows) and col_idx < len(table.rows[row_idx].cells):
                        cell = table.rows[row_idx].cells[col_idx]
                        cell_text = cell.text.strip()
                        if not cell_text:
                            cell_text = "<empty>"
                    else:
                        cell_text = "<empty>"

                    row_cells.append(cell_text)

                markdown_lines.append("| " + " | ".join(row_cells) + " |")

                if row_idx == 0:
                    separator = "| " + " | ".join(["---"] * len(row_cells)) + " |"
                    markdown_lines.append(separator)

        return '\n'.join(markdown_lines)

    def call_ai_to_analyze(self, markdown_content: str) -> Dict[str, str]:
        """调用AI分析占位符并返回替换映射"""
        print("\n🤖 正在调用AI分析占位符...")

        # 构建提示词
        prompt = f"""你是一个专业的占位符替换助手。请分析以下markdown表格和个人信息，推断每个占位符应该替换成什么内容。

**任务要求：**
1. 仔细分析表格中的占位符格式（如 {{1}}、{{2}} 等）
2. 根据个人信息推理每个占位符应该填入的内容
3. 如果无法确定某个占位符的内容，返回空字符串
4. 允许合理推理和推断
5. 返回格式必须是有效的JSON
6. JSON格式：{{"占位符编号": "要替换的内容"}}

**个人信息：**
{self.personal_info}

**Markdown表格：**
{markdown_content}

**注意：**
- 只返回需要替换的占位符
- 如果占位符内容未知或不明确，返回空字符串
- 确保JSON格式正确
- 不要在键或值中添加额外的引号或特殊字符"""

        # 调用API
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {self.api_key}"
        }

        messages = [{
            "role": "user",
            "content": prompt
        }]

        data = {
            "model": "doubao-seed-1-6-251015",
            "messages": messages,
            "thinking":{"type": "disabled"},
            "top_p": 0.7,
            "temperature": 1
        }

        try:
            print(f"🔄 正在发送请求到: https://ark.cn-beijing.volces.com/api/v3/chat/completions")
            print(f"📝 使用模型: doubao-seed-1-6-251015")
            print(f"📊 提示词长度: {len(prompt)} 字符")

            response = requests.post(
                "https://ark.cn-beijing.volces.com/api/v3/chat/completions",
                headers=headers,
                json=data,
                timeout=120
            )

            print(f"📡 响应状态码: {response.status_code}")

            if response.status_code != 200:
                print(f"❌ API调用失败!")
                print(f"响应内容: {response.text}")
                return {}

            api_response = response.json()

            if 'choices' not in api_response or not api_response['choices']:
                print(f"❌ 响应中缺少choices字段")
                return {}

            ai_response = api_response['choices'][0]['message']['content']

            print(f"\n🤖 AI响应内容：\n{ai_response}\n")

            # 解析JSON
            # 清理响应内容，移除可能的代码块标记
            json_str = ai_response.strip()
            if json_str.startswith('```json'):
                json_str = re.sub(r'^```json\s*', '', json_str)
                json_str = re.sub(r'\s*```$', '', json_str)
            elif json_str.startswith('```'):
                json_str = re.sub(r'^```\s*', '', json_str)
                json_str = re.sub(r'\s*```$', '', json_str)

            print(f"📋 清理后的JSON字符串前100字符: {json_str[:100]}...")

            # 解析JSON
            try:
                replacement_map = json.loads(json_str)
                print(f"✅ JSON解析成功")
                print(f"📊 共返回 {len(replacement_map)} 个替换项")
                return replacement_map
            except json.JSONDecodeError as e:
                print(f"❌ JSON解析失败: {e}")
                print(f"尝试解析的文本: {json_str[:200]}...")
                return {}

        except Exception as e:
            print(f"❌ AI调用失败: {e}")
            import traceback
            traceback.print_exc()
            return {}

    def replace_placeholders(self, replacement_map: Dict[str, str]):
        """根据替换映射替换docx中的占位符"""
        print("\n📝 正在替换占位符...")

        total_placeholders = 0

        # 遍历所有表格
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    original_text = cell.text
                    new_text = original_text

                    # 直接字符串替换，匹配 {1}、{2} 等格式
                    for placeholder, replacement in replacement_map.items():
                        placeholder_pattern = f"{{{placeholder}}}"
                        if placeholder_pattern in new_text:
                            new_text = new_text.replace(placeholder_pattern, replacement)

                    # 更新单元格内容
                    if new_text != original_text:
                        cell.text = new_text
                        total_placeholders += 1

        print(f"✅ 共替换了 {total_placeholders} 个占位符")

    def save(self, output_path: str):
        """保存文档"""
        self.document.save(output_path)
        print(f"\n💾 文档已保存到: {output_path}")

    def save_replacement_map(self, replacement_map: Dict[str, str], output_path: str):
        """保存替换映射到JSON文件"""
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(replacement_map, f, ensure_ascii=False, indent=2)
        print(f"📄 替换映射已保存到: {output_path}")

    def process(self, output_docx_path: str, output_json_path: str = None):
        """完整流程"""
        if output_json_path is None:
            output_json_path = output_docx_path.replace('.docx', '_replacement_map.json')

        print("=" * 60)
        print("占位符智能替换程序")
        print("=" * 60)

        # 1. 转换为markdown
        print("\n【步骤1】转换docx为markdown...")
        markdown_content = self.convert_to_markdown()
        print("✅ 转换完成")

        # 2. 保存markdown文件
        markdown_path = "带占位符表格.md"
        with open(markdown_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        print(f"📝 Markdown文件已保存到: {markdown_path}")

        # 3. 调用AI分析
        print("\n【步骤2】调用AI分析占位符...")
        replacement_map = self.call_ai_to_analyze(markdown_content)

        if not replacement_map:
            print("⚠ 没有获取到替换映射，程序终止")
            return

        # 4. 保存替换映射
        print("\n【步骤3】保存替换映射...")
        self.save_replacement_map(replacement_map, output_json_path)

        # 5. 替换占位符
        print("\n【步骤4】替换占位符...")
        self.replace_placeholders(replacement_map)

        # 6. 保存结果
        print("\n【步骤5】保存文档...")
        self.save(output_docx_path)

        print("\n" + "=" * 60)
        print("✅ 所有步骤完成！")
        print("=" * 60)


def main():
    """主函数"""
    input_file = "报名表_带占位符.docx"
    output_file = "报名表_智能替换.docx"
    personal_info_file = "个人信息.txt"

    # 检查文件是否存在
    if not os.path.exists(input_file):
        print(f"❌ 错误：找不到文件 {input_file}")
        return

    if not os.path.exists(personal_info_file):
        print(f"❌ 错误：找不到文件 {personal_info_file}")
        return

    # 获取API Key
    api_key = None

    # 1. 尝试从环境变量获取
    api_key = os.environ.get('ARK_API_KEY')

    # 2. 如果没有，尝试从api_key.txt文件读取
    if not api_key:
        api_key_file = "api_key.txt"
        if os.path.exists(api_key_file):
            with open(api_key_file, 'r', encoding='utf-8') as f:
                api_key = f.read().strip()

    # 3. 如果还是没有，使用默认值（请替换为您的实际API Key）
    if not api_key:
        api_key = "YOUR_API_KEY_HERE"  # 请替换为您的实际API Key

    if api_key == "YOUR_API_KEY_HERE":
        print("❌ 请设置您的API Key！")
        print("方法1：设置环境变量 ARK_API_KEY")
        print(f"方法2：在当前目录创建 api_key.txt 文件并写入API Key")
        print("方法3：修改代码中的api_key变量")
        return

    try:
        # 创建替换器并执行
        replacer = PlaceholderReplacer(input_file, personal_info_file, api_key)
        replacer.process(output_file)

    except Exception as e:
        print(f"\n❌ 程序执行出错: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
